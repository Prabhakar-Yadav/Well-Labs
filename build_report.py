"""
Build complete 10-page TLBC D95 Canal Blockage Detection Report
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT_PATH = r"C:\Users\PRABHAKAR\Documents\Wells-Lab\PROJECT_REPORT.docx"

doc = Document()

# ── Margins ──────────────────────────────────────────────
for sect in doc.sections:
    sect.top_margin    = Cm(1.8)
    sect.bottom_margin = Cm(1.8)
    sect.left_margin   = Cm(2.0)
    sect.right_margin  = Cm(2.0)

# ── Helper: set cell background ──────────────────────────
def bg(cell, hex_col):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_col)
    pr.append(shd)

# ── Helper: coloured paragraph ────────────────────────────
def add_para(text, size=11, bold=False, color="000000",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p

# ── Helper: table with header row ─────────────────────────
def make_table(rows, cols, header_bg="404040", header_fg="FFFFFF"):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    # header
    for cell in t.rows[0].cells:
        bg(cell, header_bg)
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor.from_string(header_fg)
            r.font.bold = True
    return t

def fill(cell, text, fg="000000", bg_col=None, size=10, bold=False,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        for r in p.runs:
            p.runs[0]._element.getparent().remove(p.runs[0]._element)
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(fg)
    if bg_col:
        bg(cell, bg_col)

# ═══════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ═══════════════════════════════════════════════════════════
add_para("TLBC D95 CANAL", 42, True, "4B0082", WD_ALIGN_PARAGRAPH.CENTER, 4)
add_para("BLOCKAGE DETECTION REPORT", 32, True, "DC143C", WD_ALIGN_PARAGRAPH.CENTER, 4)
add_para("AI-Powered Drone Survey Analysis", 20, False, "0078D7", WD_ALIGN_PARAGRAPH.CENTER, 4)
add_para("Raichur District, Karnataka", 16, False, "FF8C00", WD_ALIGN_PARAGRAPH.CENTER, 8)

# 4-stat boxes
t = doc.add_table(rows=2, cols=2)
t.style = 'Table Grid'
stats = [
    ("0", "0", "157\nCanal Reaches Analyzed",    "FF8C00"),
    ("0", "1", "87.6%\nModel Accuracy (IoU)",     "00AA00"),
    ("1", "0", "64 GB\nDrone Survey Data",         "0078D7"),
    ("1", "1", "4 Stage\nAI Pipeline",             "800080"),
]
for ri, ci, txt, col in stats:
    cell = t.cell(int(ri), int(ci))
    bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
for row in t.rows:
    row.height = Cm(1.8)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Tech bar
tb = doc.add_table(rows=1, cols=4)
tb.style = 'Table Grid'
tech = [
    ("YOLOv8 Segmentation",  "FFD080"),
    ("Mask R-CNN FIC",        "80C8FF"),
    ("DSM/DTM Silt Depth",    "90EE90"),
    ("GeoTIFF Output",        "DDA0FF"),
]
for i, (txt, col) in enumerate(tech):
    cell = tb.cell(0, i)
    bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("000000")
tb.rows[0].height = Cm(1.0)
doc.add_paragraph().paragraph_format.space_after = Pt(8)
add_para("Researcher: Prabhakar Yadav   |   prabhakaryadav0505@gmail.com   |   May 30, 2026",
         11, False, "555555", WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
add_para("What This Report Contains", 14, True, "4B0082")
rc = make_table(6, 2, "4B0082", "FFFFFF")
rcdata = [
    ("Section", "Contents"),
    ("Pages 1-2: Overview", "Project summary, study area, key findings at a glance"),
    ("Pages 3-4: Data & Pipeline", "64 GB dataset description and AI pipeline flowchart"),
    ("Pages 5-6: Methods", "Canal detection, blockage classification, FIC damage, silt depth"),
    ("Pages 7-8: Results & Notebooks", "Model performance, analysis results, notebook usage guide"),
    ("Pages 9-10: Outputs & Next Steps", "File guide, QGIS steps, action items, conclusions"),
]
for ri, (a, b) in enumerate(rcdata):
    fill(rc.cell(ri, 0), a, bold=(ri==0), bg_col="F0E8FF" if ri > 0 else None)
    fill(rc.cell(ri, 1), b, bg_col="F0E8FF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Why This Survey Matters", 14, True, "DC143C")
wm = make_table(1, 1, "FFF5E0", "000000")
fill(wm.cell(0, 0),
     "The TLBC D95 canal supplies irrigation water to thousands of acres of farmland in Raichur "
     "District. Blockages, structural damage, and silt buildup directly reduce water delivery "
     "efficiency and crop yields. This AI-powered drone survey provides the first complete, "
     "GPS-accurate health assessment of the entire canal network -- enabling field teams and "
     "engineers to prioritise targeted repairs that maximise water delivery at minimal cost.",
     "333333", "FFF5E0", 10, False)
wm.rows[0].height = Cm(1.7)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

loc = doc.add_table(rows=1, cols=3)
loc.style = 'Table Grid'
locdata = [
    ("Survey Location\nRaichur, Karnataka\nIndia", "FFE0E0"),
    ("Canal Network\nTLBC D95\n157 Reaches Analyzed", "E0F0FF"),
    ("AI Pipeline\n4 Stage Deep Learning\nMay 2026", "E0FFE8"),
]
for i, (txt, col) in enumerate(locdata):
    cell = loc.cell(0, i)
    bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("222222")
loc.rows[0].height = Cm(1.6)

# ═══════════════════════════════════════════════════════════
# PAGE 2 — PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Project Overview", 22, True, "4B0082")
add_para("The TLBC D95 canal system serves thousands of farmers in Raichur district. "
         "This project uses deep learning and drone imagery to automate inspection of the "
         "entire canal network — replacing weeks of manual surveys with an accurate, "
         "GPS-referenced AI analysis completed in days.", 11)

add_para("Three Core Problems Detected", 15, True, "DC143C")
pr = doc.add_table(rows=1, cols=3)
pr.style = 'Table Grid'
problems = [
    ("VEGETATION BLOCKAGE\n\nDense plants blocking water flow. Found in 35 reaches (22%). "
     "12 sections fully blocked.", "DC143C"),
    ("STRUCTURAL DAMAGE\n\nCracks, erosion and concrete degradation. "
     "8 critical locations flagged for repair.", "FF6400"),
    ("SILT ACCUMULATION\n\nSediment buildup measured via DSM-DTM. "
     "High silt zones (>1m) reduce canal capacity.", "0078D7"),
]
for i, (txt, col) in enumerate(problems):
    cell = pr.cell(0, i)
    bg(cell, col)
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
pr.rows[0].height = Cm(2.5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("Study Area & Survey Details", 14, True, "0078D7")
t2 = make_table(7, 2, "0078D7", "FFFFFF")
survey = [
    ("Parameter", "Value"),
    ("Location", "Raichur District, Karnataka"),
    ("Drone Platform", "DJI Phantom 4 RTK"),
    ("Ground Resolution", "~3 cm / pixel  (very high detail)"),
    ("Total Data", "64 GB  (orthomosaics + DSM + DTM)"),
    ("Coverage", "Complete D95 canal network"),
    ("Coordinate System", "WGS84 UTM Zone 43N  (EPSG:32643)"),
]
for ri, (a, b) in enumerate(survey):
    fill(t2.cell(ri, 0), a, bold=True, bg_col="DCEEFF" if ri > 0 else None)
    fill(t2.cell(ri, 1), b, bg_col="FFFFFF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("Key Findings at a Glance", 14, True, "00AA00")
kf = doc.add_table(rows=1, cols=5)
kf.style = 'Table Grid'
findings = [
    ("122\nClear Reaches",     "C8F5C8"),
    ("23\nPartial Blockage",   "FFE8B0"),
    ("12\nFully Blocked",      "FFBEBE"),
    ("8\nStructural Damage",   "FFDAB0"),
    ("1.66 GB\nOutput Files",  "D2E8FF"),
]
for i, (txt, col) in enumerate(findings):
    cell = kf.cell(0, i)
    bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(11)
    r.font.bold = True
kf.rows[0].height = Cm(1.5)

# ═══════════════════════════════════════════════════════════
# PAGE 3 — DATA DESCRIPTION
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Data Description", 22, True, "0064C8")
add_para("64 GB of high-resolution drone survey data collected over the complete TLBC D95 canal "
         "system. Three data types — orthomosaics, DSM, and DTM — feed the four AI pipeline stages.", 11)

add_para("Dataset Breakdown (64 GB Total)", 14, True, "800080")
dt = make_table(5, 4, "800080", "FFFFFF")
drows = [
    ("Data Type", "Files", "Size", "Purpose"),
    ("Orthomosaics", "5 GeoTIFF files", "59 GB",    "High-res drone stitched imagery"),
    ("DSM",          "TLBC_DSM_DTM/",   "2.4 GB",   "Surface elevation incl. silt and water"),
    ("DTM",          "TLBC_DSM_DTM/",   "2.4 GB",   "Bare earth elevation for silt calc"),
    ("Annotations",  "canal_v5 + annotation_v2", "2.78 MB", "Model training labels"),
]
row_bgs = [None, "FFDEDE", "DEE8FF", "DEFFD2", "FFFACC"]
for ri, row in enumerate(drows):
    for ci, val in enumerate(row):
        fill(dt.cell(ri, ci), val,
             bold=(ri == 0),
             bg_col=row_bgs[ri] if ri > 0 else None)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Orthomosaic Files (Kept Locally — NOT for upload due to size)", 14, True, "DC143C")
ot = make_table(6, 3, "DC143C", "FFFFFF")
orows = [
    ("File Name", "Size", "Area Covered"),
    ("TLBC_D95_ATTANUR_1_ortho.tif", "4.2 GB",  "Attanur Section 1"),
    ("TLBC_D95_ATTANUR_2_ortho.tif", "22.5 GB", "Attanur Section 2"),
    ("TLBC_D95_ATTANUR_3_ortho.tif", "17.7 GB", "Attanur Section 3"),
    ("TLBC_D95_ATTANUR_4_ortho.tif", "2.9 GB",  "Attanur Section 4"),
    ("TLBC_D95_SHAKAPUR_ortho.tif",  "10.3 GB", "Shakapur Section"),
]
for ri, row in enumerate(orows):
    for ci, val in enumerate(row):
        fill(ot.cell(ri, ci), val, bold=(ri == 0),
             bg_col="FFF0F0" if ri > 0 else None)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Training Annotations", 14, True, "00AA00")
att = make_table(3, 3, "00AA00", "FFFFFF")
arows = [
    ("Dataset", "Size", "Used For"),
    ("canal_v5", "680 KB",   "YOLOv8 canal detection training"),
    ("annotation_v2.json", "2.1 MB", "Mask R-CNN structural damage labels"),
]
for ri, row in enumerate(arows):
    for ci, val in enumerate(row):
        fill(att.cell(ri, ci), val, bold=(ri == 0),
             bg_col="EFFFEF" if ri % 2 == 1 else "FFFDE8" if ri > 0 else None)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Data Preprocessing Pipeline", 14, True, "0078D7")
pp = make_table(5, 2, "0078D7", "FFFFFF")
ppdata = [
    ("Step", "Description"),
    ("1. Tile Extraction",  "Orthomosaics split into 1024x1024 px tiles -- notebook 04_patch_tiling"),
    ("2. Canal Cropping",   "Canal regions extracted from full orthomosaic using boundary outlines -- 03_canal_extraction"),
    ("3. Annotation Load",  "canal_v5 YOLO polygon format + annotation_v2.json COCO format loaded for each model"),
    ("4. Train/Val Split",  "80% training / 20% validation split applied -- prevents overfitting"),
]
for ri, (a, b) in enumerate(ppdata):
    fill(pp.cell(ri, 0), a, bold=(ri==0), bg_col="DCF0FF" if ri > 0 else None)
    fill(pp.cell(ri, 1), b, bg_col="FFFFFF" if ri > 0 else None)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Spatial Reference System", 13, True, "800080")
cs = make_table(5, 2, "800080", "FFFFFF")
csdata = [
    ("Parameter", "Value"),
    ("Projection",  "Universal Transverse Mercator (UTM) Zone 43N"),
    ("Datum",       "WGS84 -- World Geodetic System 1984"),
    ("EPSG Code",   "EPSG:32643 -- all output GeoTIFFs use this system"),
    ("Ground Resolution", "3 cm per pixel (very high detail drone survey)"),
]
for ri, (a, b) in enumerate(csdata):
    fill(cs.cell(ri, 0), a, bold=(ri==0), bg_col="EEE4FF" if ri > 0 else None)
    fill(cs.cell(ri, 1), b, bg_col="FFFFFF" if ri > 0 else None)

# ═══════════════════════════════════════════════════════════
# PAGE 4 — PIPELINE FLOWCHART
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("AI Pipeline Architecture", 22, True, "4B0082")
add_para("Four sequential stages transform raw 64 GB drone imagery into actionable canal health maps. "
         "Each stage's output feeds directly into the next.", 11)

add_para("End-to-End Processing Flow", 15, True, "0078D7")
fc = make_table(5, 3, "1E1E1E", "FFFFFF")
fill(fc.cell(0, 0), "INPUT",             "FFFFFF", "1E1E1E", 11, True, WD_ALIGN_PARAGRAPH.CENTER)
fill(fc.cell(0, 1), "PROCESSING STAGE",  "FFFFFF", "1E1E1E", 11, True, WD_ALIGN_PARAGRAPH.CENTER)
fill(fc.cell(0, 2), "OUTPUT",            "FFFFFF", "1E1E1E", 11, True, WD_ALIGN_PARAGRAPH.CENTER)

stages = [
    ("59 GB Orthomosaics\n1024x1024 tiles",
     "STAGE 1: CANAL DETECTION\nYOLOv8 Segmentation v5\nModel: canal_detection_v5.pt\nIoU Score: 0.876 (87.6%)",
     "*_canal_mask.tif\nBinary canal boundary masks", "FFD898", "FFE8B4"),
    ("Canal mask + RGB tile",
     "STAGE 2: BLOCKAGE CLASSIFICATION\nVegetation Analysis Algorithm\nGreen Coverage + NDVI + Texture\nClasses: Clear / Partial / Blocked",
     "*_blockage.tif\nValues: 1=clear  2=partial  3=blocked", "B4E8B4", "D2F5D2"),
    ("Canal regions + annotation_v2.json",
     "STAGE 3: FIC DAMAGE DETECTION\nMask R-CNN + ResNet50\nModel: fic_detector.pth\nTrain: 0.23  Val: 0.18",
     "canal_data.json\nstructural_damage: true/false", "FFD0A0", "FFE4C4"),
    ("DSM (2.4 GB) + DTM (2.4 GB)",
     "STAGE 4: SILT DEPTH MEASUREMENT\nElevation Differencing\nFormula: Silt = DSM - DTM\nClasses: Low / Medium / High",
     "*_silt.tif + *_silt_depth.tif\nDepth in metres", "C0D8FF", "DCECFF"),
]
for ri, (inp, proc, out, left_bg, right_bg) in enumerate(stages):
    fill(fc.cell(ri+1, 0), inp,  bg_col=left_bg,  size=9)
    fill(fc.cell(ri+1, 1), proc, bg_col=left_bg,  size=9, bold=True)
    fill(fc.cell(ri+1, 2), out,  bg_col=right_bg, size=9)
    for col in range(3):
        from docx.shared import Cm as cm2
        fc.cell(ri+1, col)._tc.get_or_add_tcPr()
for row in fc.rows[1:]:
    row.height = Cm(1.7)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Stages run sequentially. Stage 1 output feeds Stage 2, and so on. All outputs are "
         "georeferenced GeoTIFFs in WGS84 UTM Zone 43N ready for QGIS.", 9, color="666666")

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("Processing Summary", 14, True, "0078D7")
ps = doc.add_table(rows=1, cols=4)
ps.style = 'Table Grid'
psum = [
    ("64 GB\nRaw Input Data",      "DCDCFF"),
    ("7\nNotebooks Run",           "FFE8CC"),
    ("157\nReaches Analyzed",      "CCFFCC"),
    ("1.66 GB\nFinal Outputs",     "FFCCCC"),
]
for i, (txt, col) in enumerate(psum):
    cell = ps.cell(0, i)
    bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(11)
    r.font.bold = True
ps.rows[0].height = Cm(1.4)

# ═══════════════════════════════════════════════════════════
# PAGE 5 — STAGE 1 & 2
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Methodology: Canal Detection and Blockage", 22, True, "0078D7")

add_para("STAGE 1: Canal Boundary Detection", 16, True, "FF8C00")
add_para("YOLOv8 Segmentation v5 was trained on 680 KB of annotated canal polygons to identify "
         "precise canal boundaries in 1024x1024 px image tiles. The model outputs pixel-level masks "
         "that are fully georeferenced and align with real GPS coordinates.", 11)
s1 = make_table(7, 2, "FF8C00", "FFFFFF")
s1data = [
    ("Field", "Value"),
    ("Model Architecture", "YOLOv8 Segmentation Version 5"),
    ("Training Data", "canal_v5  --  680 KB of polygon annotations"),
    ("Input Tile Size", "1024 x 1024 pixels"),
    ("IoU Score", "0.876  (87.6% overlap with ground truth -- Excellent)"),
    ("Output File", "*_canal_mask.tif  --  binary canal boundary raster"),
    ("What IoU means", "Predictions match expert annotations 87.6% -- production ready"),
]
for ri, (a, b) in enumerate(s1data):
    fill(s1.cell(ri, 0), a, bold=(ri==0), bg_col="FFF0D8" if ri > 0 else None)
    fill(s1.cell(ri, 1), b, bg_col="FFFFFF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("STAGE 2: Blockage Classification", 16, True, "00A000")
add_para("A rule-based algorithm analyses vegetation within each detected canal segment using "
         "three spectral metrics. Results are explicit and explainable -- no black box.", 11)
s2 = make_table(4, 3, "00A000", "FFFFFF")
s2data = [
    ("Metric", "What It Measures", "Thresholds"),
    ("Green Coverage %",  "Vegetation vs open water surface",    "Less than 20 = Clear | 20-50 = Partial | Over 50 = Blocked"),
    ("NDVI Index",        "Vegetation density and health",       "Below 0.3 = Clear | 0.3-0.5 = Partial | Above 0.5 = Blocked"),
    ("Texture Analysis",  "Rough plants vs smooth water",        "Low texture = clear | High texture = vegetation"),
]
for ri, row in enumerate(s2data):
    for ci, val in enumerate(row):
        fill(s2.cell(ri, ci), val, bold=(ri==0), bg_col="EFFFEF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("Blockage Classification Legend", 13, True, "000000")
cl = make_table(4, 4, "303030", "FFFFFF")
cldata = [
    ("Class", "Color in QGIS", "Vegetation", "Required Action"),
    ("1 -- Clear",   "Green",  "Less than 20%", "No action needed"),
    ("2 -- Partial", "Orange", "20 to 50%",     "Clear within 2-3 months"),
    ("3 -- Blocked", "Red",    "Over 50%",       "URGENT -- clear immediately"),
]
clbgs = [None, "C8FFC8", "FFE0A0", "FFBEBE"]
for ri, row in enumerate(cldata):
    for ci, val in enumerate(row):
        fill(cl.cell(ri, ci), val, bold=(ri==0), bg_col=clbgs[ri] if ri > 0 else None,
             align=WD_ALIGN_PARAGRAPH.CENTER)

# ═══════════════════════════════════════════════════════════
# PAGE 6 — STAGE 3 & 4
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Methodology: Damage Detection and Silt Measurement", 22, True, "800080")

add_para("STAGE 3: Structural Damage Detection (FIC)", 16, True, "FF6400")
add_para("Mask R-CNN with ResNet50 backbone detects cracks, erosion, and concrete degradation. "
         "Trained on 2.1 MB of expert-labeled FIC annotations. Detection threshold set at 0.2 "
         "for higher sensitivity -- better to flag for review than miss real structural damage.", 11)
s3 = make_table(8, 2, "FF6400", "FFFFFF")
s3data = [
    ("Field", "Value"),
    ("Architecture",        "Mask R-CNN + ResNet50 backbone"),
    ("Training Data",       "annotation_v2.json  --  2.1 MB FIC damage labels"),
    ("Detection Threshold", "0.2  (lowered from default 0.4 for higher sensitivity)"),
    ("Morph. Filter",       "3x3 kernel, min area 20 pixels"),
    ("Training Loss",       "0.23  --  model converged well"),
    ("Validation Loss",     "0.18  --  lower than train loss, no overfitting"),
    ("Damage Locations",    "8 reaches flagged with structural_damage = true in canal_data.json"),
]
for ri, (a, b) in enumerate(s3data):
    fill(s3.cell(ri, 0), a, bold=(ri==0), bg_col="FFF0E4" if ri > 0 else None)
    fill(s3.cell(ri, 1), b, bg_col="FFFFFF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("STAGE 4: Silt Depth Measurement", 16, True, "0096C8")
add_para("Silt is measured by subtracting bare-earth elevation (DTM) from surface elevation (DSM). "
         "Where DSM > DTM inside canal boundaries, the difference equals silt depth in metres. "
         "Only canal pixels (from Stage 1 mask) are included to avoid false positives.", 11)

# Formula highlight box
fh = make_table(1, 1, "C8E8FF", "4B0082")
fill(fh.cell(0, 0), "FORMULA:      Silt Depth (metres)  =  DSM  -  DTM",
     "4B0082", "C8E8FF", 14, True, WD_ALIGN_PARAGRAPH.CENTER)
fh.rows[0].height = Cm(1.2)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

s4 = make_table(4, 4, "0096C8", "FFFFFF")
s4data = [
    ("Class", "Depth Range", "QGIS Colour", "Action Required"),
    ("Low (1)",    "Less than 0.5 m", "Light Yellow", "Monitor only"),
    ("Medium (2)", "0.5 m to 1.0 m", "Orange",        "Schedule dredging within 3 months"),
    ("High (3)",   "Greater than 1.0 m", "Dark Red",  "Immediate dredging required"),
]
s4bgs = [None, "FFFF88", "FFD080", "FFB4B4"]
for ri, row in enumerate(s4data):
    for ci, val in enumerate(row):
        fill(s4.cell(ri, ci), val, bold=(ri==0),
             bg_col=s4bgs[ri] if ri > 0 else None,
             align=WD_ALIGN_PARAGRAPH.CENTER)

# ═══════════════════════════════════════════════════════════
# PAGE 7 — RESULTS
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Results and Model Performance", 22, True, "DC143C")

add_para("Model Accuracy Summary", 15, True, "800080")
mp = make_table(4, 4, "800080", "FFFFFF")
mpdata = [
    ("Model", "Metric", "Score", "Assessment"),
    ("Canal Detection v5", "IoU Score",   "0.876", "Excellent -- production ready"),
    ("FIC Detector",        "Train Loss",  "0.23",  "Converged well"),
    ("FIC Detector",        "Val Loss",    "0.18",  "No overfitting -- generalises well"),
]
mpbgs = [None, "C8FFC8", "DCF0FF", "FFF0DC"]
for ri, row in enumerate(mpdata):
    for ci, val in enumerate(row):
        fill(mp.cell(ri, ci), val, bold=(ri==0), bg_col=mpbgs[ri] if ri > 0 else None,
             align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Blockage Analysis -- 157 Canal Reaches", 15, True, "DC143C")
br = make_table(6, 4, "1E1E1E", "FFFFFF")
brdata = [
    ("Status", "Count", "Percentage", "Required Action"),
    ("TOTAL",            "157", "100%",   "--"),
    ("Clear",            "122", "77.7%",  "No action needed"),
    ("Partial Blockage", "23",  "14.6%",  "Clear within 2-3 months"),
    ("Fully Blocked",    "12",  "7.6%",   "URGENT -- clear immediately"),
    ("Structural Damage","8",   "5.1%",   "Engineering inspection required"),
]
brbgs = [None, "E8E8E8", "C8FFC8", "FFE0A0", "FFBEBE", "FFD0A0"]
for ri, row in enumerate(brdata):
    for ci, val in enumerate(row):
        fill(br.cell(ri, ci), val, bold=(ri == 0 or ci == 0),
             bg_col=brbgs[ri] if ri > 0 else None,
             align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Structural Damage: 8 Locations Detected", 14, True, "FF6400")
add_para("Canal reaches flagged by FIC detector show cracks, erosion, or concrete deterioration. "
         "Load canal_data.json and filter for structural_damage = true to get GPS coordinates for each. "
         "These locations need engineering assessment before further deterioration causes canal failure.", 11)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Impact box
ib = make_table(1, 1, "E0FFE0", "000000")
fill(ib.cell(0,0),
     "PROJECT IMPACT:  This automated system replaces weeks of manual surveys. "
     "Field teams now have GPS-accurate locations for all 35 problem reaches. "
     "Engineers have precise coordinates for 8 repair sites. Analysis completed in days, not weeks.",
     "005500", "E0FFE0", 11, True)
ib.rows[0].height = Cm(1.5)

# ═══════════════════════════════════════════════════════════
# PAGE 8 — NOTEBOOKS GUIDE
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Pipeline Notebooks", 22, True, "4B0082")
add_para("7 Jupyter notebooks implement the full pipeline. Must be run in the exact order shown -- "
         "each notebook's output feeds the next stage.", 11)

nb = make_table(8, 3, "4B0082", "FFFFFF")
nbdata = [
    ("Notebook", "Purpose", "Output"),
    ("03_canal_extraction",   "Crops canal regions from large orthomosaics",            "Focused canal tiles"),
    ("04_dataset",            "Builds training dataset from canal_v5 annotations",       "Train/val image + mask pairs"),
    ("04_patch_tiling",       "Splits images into 1024x1024 px tiles for models",        "Thousands of image tiles"),
    ("05_binary_training",    "Trains YOLOv8 canal detector",                            "canal_detection_v5.pt  IoU=0.876"),
    ("06_fic_training",       "Trains Mask R-CNN damage detection model",               "fic_detector.pth  ValLoss=0.18"),
    ("07_blockage_detection", "Runs full inference and classification",                  "1.6 GB blockage GeoTIFFs"),
    ("08_silt_detection",     "DSM minus DTM silt depth measurement",                    "57 MB silt GeoTIFFs"),
]
nbbgs = [None, "E8F0FF", "F0FFE8", "FFF8E8", "E8FFE8", "FFE8F0", "FFECEC", "E8F4FF"]
for ri, row in enumerate(nbdata):
    for ci, val in enumerate(row):
        fill(nb.cell(ri, ci), val, bold=(ri==0),
             bg_col=nbbgs[ri] if ri > 0 else None, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Critical warning
wb = make_table(1, 1, "DC143C", "FFFFFF")
fill(wb.cell(0, 0),
     "CRITICAL -- Notebook 07 Cell Execution Order\n\n"
     "Cell 8 (visualization) MUST run AFTER: Cell 1 -> Cell 2b -> Cell 3 -> Cell 4\n"
     "Reason: Cell 3 creates 'results' dict | Cell 4 adds 'status' + 'structural_damage'\n"
     "Running Cell 8 alone will cause NameError -- the variables do not exist yet!",
     "FFFFFF", "DC143C", 11, True)
wb.rows[0].height = Cm(2.2)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Notebook 07 -- Cell Dependencies", 13, True, "DC143C")
cb = make_table(5, 2, "DC143C", "FFFFFF")
cbdata = [
    ("Cell", "What It Creates"),
    ("Cell 1",  "Loads canal_detection_v5.pt model into memory"),
    ("Cell 2b", "Builds canal_data dictionary with reach geometries"),
    ("Cell 3",  "Runs YOLOv8 inference -- creates 'results' dictionary"),
    ("Cell 4",  "Runs FIC detector -- adds 'status' and 'structural_damage' fields"),
]
for ri, (a, b) in enumerate(cbdata):
    fill(cb.cell(ri, 0), a, bold=(ri==0), bg_col="FFF0F0" if ri > 0 else None)
    fill(cb.cell(ri, 1), b, bg_col="FFF0F0" if ri > 0 else None)

# ═══════════════════════════════════════════════════════════
# PAGE 9 — OUTPUT FILES GUIDE
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Output Files Guide", 22, True, "FF6400")

add_para("Blockage Outputs  --  outputs/07_blockage/  [1.6 GB]", 15, True, "DC143C")
of = make_table(4, 3, "DC143C", "FFFFFF")
ofdata = [
    ("File", "Contains", "How to Use"),
    ("*_blockage.tif",
     "Pixel values: 0=background  1=clear  2=partial  3=blocked",
     "QGIS: Symbology > Paletted. Set 1=green, 2=orange, 3=red. Overlay on ortho at 60% opacity."),
    ("*_canal_mask.tif",
     "Binary: 1=canal detected  0=not canal",
     "Verify detection accuracy. Export as shapefile for precise boundary mapping."),
    ("canal_data.json",
     "reach_id, status, structural_damage flag, GPS polygon coords for all 157 reaches",
     "Filter status='blocked' or structural_damage=true to generate field team task list."),
]
ofbgs = [None, "FFEAEA", "EAF0FF", "EAFFEA"]
for ri, row in enumerate(ofdata):
    for ci, val in enumerate(row):
        fill(of.cell(ri, ci), val, bold=(ri==0), bg_col=ofbgs[ri] if ri > 0 else None, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Silt Outputs  --  outputs/08_silt/  [57 MB]", 15, True, "0096C8")
sf = make_table(3, 3, "0096C8", "FFFFFF")
sfdata = [
    ("File", "Contains", "How to Use"),
    ("*_silt.tif",
     "Classified: 1=low less than 0.5m  2=medium 0.5-1m  3=high over 1m",
     "QGIS: Apply yellow-orange-red gradient. Red zones need immediate dredging."),
    ("*_silt_depth.tif",
     "Continuous depth values in metres (raw)",
     "Calculate dredging volume. Track accumulation over repeat surveys. Cost estimation."),
]
sfbgs = [None, "E0FFFF", "E0F0FF"]
for ri, row in enumerate(sfdata):
    for ci, val in enumerate(row):
        fill(sf.cell(ri, ci), val, bold=(ri==0), bg_col=sfbgs[ri] if ri > 0 else None, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Upload Status", 15, True, "00AA00")
ul = make_table(3, 3, "303030", "FFFFFF")
uldata = [
    ("Category", "Contents", "Size"),
    ("UPLOAD TO DRIVE (share with team)",
     "outputs/ + all notebooks + annotation/ + annotation_v2/",
     "4.5 GB"),
    ("KEEP LOCALLY (_keep_locally/)",
     "Orthomosaics + DSM/DTM + manual_labels + .claude/",
     "59.5 GB"),
]
ulbgs = [None, "CCFFCC", "FFEECC"]
for ri, row in enumerate(uldata):
    for ci, val in enumerate(row):
        fill(ul.cell(ri, ci), val, bold=(ri==0), bg_col=ulbgs[ri] if ri > 0 else None)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_para("QGIS Colour Rendering Reference", 13, True, "800080")
qr = make_table(7, 3, "800080", "FFFFFF")
qrdata = [
    ("File Type",       "Pixel Value",            "QGIS Colour / Action"),
    ("*_blockage.tif",  "1 = Clear Canal",         "Green #00AA00 -- no action needed"),
    ("*_blockage.tif",  "2 = Partial Blockage",    "Orange #FF8C00 -- clear within 3 months"),
    ("*_blockage.tif",  "3 = Fully Blocked",       "Red #DC143C -- URGENT immediate clearing"),
    ("*_silt.tif",      "1 = Low (less 0.5m)",     "Light Yellow #FFFACD -- monitor only"),
    ("*_silt.tif",      "2 = Medium (0.5 to 1m)",  "Orange #FFD700 -- schedule dredging"),
    ("*_silt.tif",      "3 = High (over 1m)",      "Dark Red #8B0000 -- immediate dredging"),
]
qrbgs = [None, "C8FFC8", "FFE0A0", "FFBEBE", "FFFACD", "FFD080", "FFB4B4"]
for ri, row in enumerate(qrdata):
    for ci, val in enumerate(row):
        fill(qr.cell(ri, ci), val, bold=(ri==0), bg_col=qrbgs[ri] if ri > 0 else None, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_para("File Naming Pattern", 14, True, "FF6400")
fn = make_table(5, 2, "FF6400", "FFFFFF")
fndata = [
    ("Pattern", "Example and Meaning"),
    ("TLBC_D95_{AREA}_{N}_ortho.tif",  "TLBC_D95_ATTANUR_1_ortho.tif -- raw orthomosaic input"),
    ("{AREA}_{N}_canal_mask.tif",       "ATTANUR_1_canal_mask.tif -- binary canal boundary raster"),
    ("{AREA}_{N}_blockage.tif",         "ATTANUR_1_blockage.tif -- 3-class blockage raster"),
    ("{AREA}_{N}_silt.tif",             "ATTANUR_1_silt.tif -- 3-class silt depth raster"),
]
fnbgs = [None, "FFF0E0", "FFE8CC", "FFDCB4", "FFCCAA"]
for ri, (a, b) in enumerate(fndata):
    fill(fn.cell(ri, 0), a, bold=(ri==0), bg_col=fnbgs[ri] if ri > 0 else None)
    fill(fn.cell(ri, 1), b, bg_col=fnbgs[ri] if ri > 0 else None)

# ═══════════════════════════════════════════════════════════
# PAGE 10 — CONCLUSION & NEXT STEPS
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para("Conclusion and Next Steps", 22, True, "00845A")
add_para("This project deployed a four-stage AI pipeline to inspect the entire TLBC D95 canal network. "
         "Processing 64 GB of drone data, the system identified all blockage zones, structural damage "
         "sites, and silt accumulation areas with GPS accuracy -- results are immediately deployable "
         "by field teams and engineers without further analysis.", 11)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_para("Deliverables Summary", 15, True, "0078D7")
dl = make_table(5, 1, "0078D7", "FFFFFF")
dldata = [
    ("Deliverable"),
    ("Canal detection model: 87.6% IoU -- production-ready for new drone surveys"),
    ("Damage detection model: Val loss 0.18 -- generalises well to unseen canal sections"),
    ("1.6 GB blockage GeoTIFFs: pixel-level classification of all 157 reaches"),
    ("57 MB silt GeoTIFFs: depth classification + continuous values for dredging estimates"),
]
dlbgs = [None, "CCFFCC", "CCE8FF", "FFEECC", "FFDCCC"]
for ri, row in enumerate(dldata):
    fill(dl.cell(ri, 0), row, bold=(ri==0), bg_col=dlbgs[ri] if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("Priority Action Items", 15, True, "DC143C")
ai = make_table(4, 3, "1E1E1E", "FFFFFF")
aidata = [
    ("Priority", "Action", "Deadline"),
    ("HIGH",   "Deploy field teams to 12 blocked reaches -- clear vegetation", "Immediate"),
    ("HIGH",   "Engineering assessment of 8 structural damage locations",       "Within 2 weeks"),
    ("MEDIUM", "Dredge high silt zones (over 1m depth) to restore capacity",   "Within 3 months"),
]
aibgs = [None, "FFCCCC", "FFD4B0", "FFFACC"]
for ri, row in enumerate(aidata):
    for ci, val in enumerate(row):
        fill(ai.cell(ri, ci), val, bold=(ri==0 or ci==0),
             bg_col=aibgs[ri] if ri > 0 else None, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para("QGIS Quick-Start", 15, True, "800080")
qs = make_table(5, 2, "800080", "FFFFFF")
qsdata = [
    ("Step", "Action"),
    ("1", "Layer > Add Raster > load orthomosaic as base reference layer"),
    ("2", "Layer > Add Raster > load *_blockage.tif as overlay"),
    ("3", "Symbology: Paletted/Unique Values -- set 1=green, 2=orange, 3=red, 0=transparent"),
    ("4", "Set layer transparency to 60% -- see ortho and classification together"),
]
for ri, (a, b) in enumerate(qsdata):
    fill(qs.cell(ri, 0), a, bold=(ri==0), bg_col="F0E8FF" if ri > 0 else None,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    fill(qs.cell(ri, 1), b, bg_col="F0E8FF" if ri > 0 else None)
doc.add_paragraph().paragraph_format.space_after = Pt(10)

add_para("Prabhakar Yadav   |   prabhakaryadav0505@gmail.com   |   May 30, 2026",
         12, True, "0078D7", WD_ALIGN_PARAGRAPH.CENTER)
add_para("TLBC D95 Canal Blockage Detection -- AI-Powered Infrastructure Monitoring, Raichur District",
         10, False, "888888", WD_ALIGN_PARAGRAPH.CENTER)

# ── SAVE ─────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"SAVED: {OUT_PATH}")
print(f"Pages: 10  |  All sections complete")
